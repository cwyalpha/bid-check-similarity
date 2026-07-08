from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image

from checksim.engine import _hamming_hex, load_config, run_check
from checksim.parsers import _paddleocr_result_blocks, parse_file
from checksim.models import CheckOptions
from checksim.text import split_blocks_to_units


class CheckSimEngineTest(unittest.TestCase):
    def test_core_detection_exclusion_keywords_images_and_offline_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image_path = root / "logo.png"
            Image.new("RGB", (16, 16), color=(20, 120, 90)).save(image_path)

            docx_path = root / "A公司投标文件.docx"
            _make_docx(
                docx_path,
                [
                    "目录",
                    "本项目采用统一身份认证、安全审计和日志留存方案，确保业务系统在内网环境中稳定运行并形成可追溯记录。",
                    "投标人应完全响应招标文件中的服务范围、验收标准和付款条件。",
                    "杭州天衡科技有限公司负责本项目实施。",
                ],
                image_path,
            )

            md_path = root / "B公司投标文件.md"
            md_path.write_text(
                "\n".join(
                    [
                        "# 目录",
                        "本项目采用统一身份认证、安全审计和日志留存方案，确保业务系统在内网环境中稳定运行并形成可追溯记录。",
                        "投标人应完全响应招标文件中的服务范围、验收标准和付款条件。",
                        "杭州天衡科技有限公司作为合作单位参与本项目。",
                        "![logo](logo.png)",
                    ]
                ),
                encoding="utf-8",
            )

            exclude_path = root / "招标文件.md"
            exclude_path.write_text("投标人应完全响应招标文件中的服务范围、验收标准和付款条件。", encoding="utf-8")

            config = {
                "groups": [
                    {"name": "A公司", "files": [str(docx_path)]},
                    {"name": "B公司", "files": [str(md_path)]},
                ],
                "exclude_files": [str(exclude_path)],
                "keywords": ["杭州天衡科技有限公司", "re:统一身份认证"],
                "options": {
                    "min_chars": 20,
                    "min_words": 8,
                    "similarity_threshold": 0.78,
                    "exclude_threshold": 0.86,
                },
            }

            result = run_check(config, root / "outputs")
            active = [match for match in result["matches"] if not match["excluded"]]
            excluded = [match for match in result["matches"] if match["excluded"]]

            self.assertGreaterEqual(len(active), 1)
            self.assertTrue(any("统一身份认证" in match["text_a"] + match["text_b"] for match in active))
            self.assertTrue(any("服务范围" in match["text_a"] + match["text_b"] for match in excluded))
            self.assertFalse(any(match["text_a"] == "目录" or match["text_b"] == "目录" for match in active))
            self.assertGreaterEqual(result["stats"]["short_filtered_unit_count"], 1)
            self.assertGreaterEqual(result["stats"]["keyword_alert_count"], 2)
            self.assertGreaterEqual(result["stats"]["image_duplicate_count"], 1)
            self.assertTrue(all("unit_id_a" in match and "unit_id_b" in match for match in result["matches"]))
            all_units = {unit["unit_id"] for doc in result["documents"] for unit in doc["units"]}
            self.assertTrue(all(match["unit_id_a"] in all_units and match["unit_id_b"] in all_units for match in result["matches"]))
            self.assertTrue(result["output_files"]["compare_pages"])
            self.assertIn("ai_summary_json", result["output_files"])
            self.assertGreater(result["output_files"]["ai_summary_json_bytes"], 0)

            report_html = Path(result["output_files"]["report_html"]).read_text(encoding="utf-8")
            self.assertIn("<style>", report_html)
            self.assertIn("<script>", report_html)
            self.assertIn("左右对照", report_html)
            self.assertIn("AI 精简结果", report_html)
            self.assertNotIn("cdn.jsdelivr", report_html)
            self.assertNotIn("https://", report_html)

            ai_summary = json.loads(Path(result["output_files"]["ai_summary_json"]).read_text(encoding="utf-8"))
            self.assertEqual(ai_summary["schema"], "checksim.ai_summary.v1")
            self.assertEqual(ai_summary["output_files"]["report_html"], result["output_files"]["report_html"])
            self.assertEqual(ai_summary["output_files"]["result_json"], result["output_files"]["result_json"])
            self.assertGreaterEqual(ai_summary["evidence"]["similar_text"]["sample_count"], 1)
            self.assertGreaterEqual(ai_summary["evidence"]["keyword_alerts"][0]["hit_count"], 2)
            self.assertNotIn("documents", ai_summary)

            compare_path = Path(result["output_files"]["compare_pages"][0]["path"])
            compare_html = compare_path.read_text(encoding="utf-8")
            self.assertIn("返回总报告", compare_html)
            self.assertIn("marker-rail", compare_html)
            self.assertIn("上一个高亮", compare_html)
            self.assertIn("下一个高亮", compare_html)
            self.assertIn("包含已排除", compare_html)
            self.assertIn("doc-stream", compare_html)
            self.assertNotIn("https://", compare_html)
            self.assertNotIn("<script src", compare_html)

    def test_docx_tables_are_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "table.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "服务期限"
            table.cell(0, 1).text = "三年驻场运维服务"
            doc.save(path)

            parsed = parse_file(path, "测试组", 0, CheckOptions())
            self.assertTrue(any("三年驻场运维服务" in block for block in parsed.blocks))

    def test_load_config_resolves_paths_relative_to_config_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            case_dir = root / "case"
            case_dir.mkdir()
            (case_dir / "A公司").mkdir()
            (case_dir / "B公司").mkdir()
            (case_dir / "招标文件.docx").write_bytes(b"placeholder")
            config_path = case_dir / "case.json"
            config_path.write_text(
                json.dumps(
                    {
                        "groups": [
                            {"name": "A公司", "files": ["A公司/投标文件.docx"]},
                            {"name": "B公司", "files": ["B公司/投标文件.docx"]},
                        ],
                        "exclude_files": ["招标文件.docx"],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            config = load_config(config_path)

        self.assertEqual(config["groups"][0]["files"], [str((case_dir / "A公司/投标文件.docx").resolve())])
        self.assertEqual(config["groups"][1]["files"], [str((case_dir / "B公司/投标文件.docx").resolve())])
        self.assertEqual(config["exclude_files"], [str((case_dir / "招标文件.docx").resolve())])

    def test_legacy_doc_is_converted_then_parsed_with_original_name(self) -> None:
        temp_dirs: list[Path] = []

        def fake_convert(path: str | Path, **_kwargs: object) -> tuple[Path, Path]:
            temp_dir = Path(tempfile.mkdtemp())
            temp_dirs.append(temp_dir)
            converted = temp_dir / "converted.docx"
            _make_docx(converted, ["旧格式投标文件正文，转换后应进入查重解析流程。"], None)
            return converted, temp_dir

        with tempfile.TemporaryDirectory() as tmp:
            legacy_path = Path(tmp) / "legacy.doc"
            legacy_path.write_bytes(b"fake legacy content")
            with patch("checksim.parsers.convert_legacy_to_docx", side_effect=fake_convert) as mocked:
                parsed = parse_file(legacy_path, "测试组", 0, CheckOptions(min_chars=1))

            mocked.assert_called_once()
            self.assertEqual(parsed.file_name, "legacy.doc")
            self.assertEqual(Path(parsed.file_path), legacy_path.resolve())
            self.assertEqual(parsed.metadata["converted_from"], ".doc")
            self.assertTrue(any("旧格式投标文件正文" in block for block in parsed.blocks))
            self.assertTrue(temp_dirs)
            self.assertFalse(temp_dirs[0].exists())

    def test_near_duplicate_above_threshold_is_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            a = root / "a.md"
            b = root / "b.md"
            a.write_text(
                "本项目采用统一身份认证、安全审计和日志留存方案，确保业务系统在内网环境中稳定运行并形成可追溯记录；系统上线后将建立每日巡检、异常告警和月度运行分析机制。",
                encoding="utf-8",
            )
            b.write_text(
                "本项目采用统一身份认证、安全审计及日志留存方案，保障业务系统在内网环境中稳定运行并形成可追溯记录；上线后建立每日巡检、异常预警和月度运行分析机制。",
                encoding="utf-8",
            )
            result = run_check(
                {
                    "groups": [{"name": "甲公司", "files": [str(a)]}, {"name": "乙公司", "files": [str(b)]}],
                    "options": {"similarity_threshold": 0.78, "sentence_delimiters": "。！？!?；;"},
                }
            )
            scores = [match["similarity"] for match in result["matches"] if not match["excluded"]]
            self.assertTrue(scores)
            self.assertGreater(max(scores), 0.78)
            self.assertLess(max(scores), 1.0)

    def test_sentence_delimiters_are_configurable(self) -> None:
        text = "第一句需要单独比较；第二句也需要单独比较。第三句保留。"
        with_semicolon = split_blocks_to_units([text], 420, "。；", "，")
        without_semicolon = split_blocks_to_units([text], 420, "。", "，")

        self.assertEqual(len(with_semicolon), 3)
        self.assertEqual(len(without_semicolon), 2)

    def test_hamming_hex_is_python38_compatible(self) -> None:
        self.assertEqual(_hamming_hex("0000000000000000", "000000000000000f"), 4)

    def test_compare_page_marks_one_to_many_matches(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            left = root / "left.md"
            right = root / "right.md"
            left.write_text(
                "本系统采用统一身份认证、安全审计和日志留存方案，确保业务系统在内网环境中稳定运行并形成可追溯记录。",
                encoding="utf-8",
            )
            right.write_text(
                "\n".join(
                    [
                        "本系统采用统一身份认证、安全审计及日志留存方案，保障业务系统在内网环境中稳定运行并形成可追溯记录。",
                        "本系统将采用统一身份认证、安全审计和日志留存机制，确保业务系统在内网环境中稳定运行并形成追溯记录。",
                    ]
                ),
                encoding="utf-8",
            )
            result = run_check(
                {
                    "groups": [{"name": "甲公司", "files": [str(left)]}, {"name": "乙公司", "files": [str(right)]}],
                    "options": {"similarity_threshold": 0.78},
                },
                root / "outputs",
            )
            active = [match for match in result["matches"] if not match["excluded"]]
            self.assertGreaterEqual(len(active), 2)
            compare_html = Path(result["output_files"]["compare_pages"][0]["path"]).read_text(encoding="utf-8")
            self.assertIn("hit-badge", compare_html)
            self.assertIn(">2</span>", compare_html)
            self.assertIn("data-targets=", compare_html)

    def test_markdown_markup_is_cleaned_from_report_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.md"
            path.write_text(
                "\n".join(
                    [
                        "# 项目标题",
                        "**加粗文本** 和 _斜体文本_",
                        "- 列表项",
                        "1. 有序项",
                        "> 引用内容",
                        "| 列 | 值 |",
                        "| --- | --- |",
                        "| 服务 | 内容 |",
                        "[链接文字](https://example.invalid)",
                    ]
                ),
                encoding="utf-8",
            )
            parsed = parse_file(path, "测试组", 0, CheckOptions(min_chars=1))
            text = "\n".join(parsed.blocks)
            self.assertIn("项目标题", text)
            self.assertIn("加粗文本 和 斜体文本", text)
            self.assertIn("列表项", text)
            self.assertIn("有序项", text)
            self.assertIn("引用内容", text)
            self.assertIn("链接文字", text)
            self.assertNotIn("#", text)
            self.assertNotIn("**", text)
            self.assertNotIn("|", text)

    def test_result_json_keeps_full_matches_without_truncation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            left = root / "left.md"
            right = root / "right.md"
            base_lines = [
                f"第{index}段 本项目采用统一身份认证、安全审计、日志留存和月度运行分析机制，确保业务系统稳定运行并形成可追溯记录。"
                for index in range(1, 9)
            ]
            left.write_text("\n\n".join(base_lines), encoding="utf-8")
            right.write_text(
                "\n\n".join(line.replace("确保", "保障").replace("可追溯记录", "追溯记录") for line in base_lines),
                encoding="utf-8",
            )

            result = run_check(
                {
                    "groups": [{"name": "甲公司", "files": [str(left)]}, {"name": "乙公司", "files": [str(right)]}],
                    "options": {
                        "min_chars": 10,
                        "similarity_threshold": 0.78,
                    },
                },
                root / "outputs",
            )

            stats = result["stats"]
            self.assertFalse(stats["match_truncated"])
            self.assertEqual(stats["displayed_similar_match_count"], stats["total_similar_match_count"])
            self.assertEqual(len([match for match in result["matches"] if not match["excluded"]]), stats["total_similar_match_count"])

            output_files = result["output_files"]
            self.assertNotIn("all_matches_jsonl", output_files)
            self.assertTrue(Path(output_files["ai_summary_json"]).exists())

            result_json = Path(output_files["result_json"]).read_text(encoding="utf-8")
            self.assertNotIn('"_all_matches"', result_json)
            ai_summary = json.loads(Path(output_files["ai_summary_json"]).read_text(encoding="utf-8"))
            self.assertEqual(ai_summary["stats"]["similar_match_count"], stats["similar_match_count"])
            self.assertLess(Path(output_files["ai_summary_json"]).stat().st_size, Path(output_files["result_json"]).stat().st_size)
            report_html = Path(output_files["report_html"]).read_text(encoding="utf-8")
            self.assertNotIn("已截断", report_html)
            self.assertNotIn("all_matches.jsonl", report_html)
            self.assertIn("pairSummaryTable", report_html)
            self.assertIn("minSimFilter", report_html)
            self.assertIn("pairIncludeExcluded", report_html)

            compare_html = Path(output_files["compare_pages"][0]["path"]).read_text(encoding="utf-8")
            self.assertNotIn("代表性", compare_html)

    def test_similarity_backend_embedding_is_reserved(self) -> None:
        with self.assertRaisesRegex(ValueError, "embedding"):
            CheckOptions.from_dict({"similarity_backend": "embedding"})

    def test_txt_is_parsed_as_plain_text_without_markdown_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "plain.txt"
            path.write_text("# 标题\n**这不是 Markdown 加粗标记**\n普通纯文本内容。", encoding="utf-8")

            parsed = parse_file(path, "测试组", 0, CheckOptions(min_chars=1))
            text = "\n".join(parsed.blocks)

            self.assertIn("# 标题", text)
            self.assertIn("**这不是 Markdown 加粗标记**", text)
            self.assertTrue(parsed.units)

    def test_pdf_text_layer_is_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.pdf"
            _make_pdf(path, "Alpha Company tender PDF text for procurement audit.")

            parsed = parse_file(path, "PDF组", 0, CheckOptions(min_chars=1, pdf_ocr_mode="off"))
            text = "\n".join(parsed.blocks)

            self.assertIn("Alpha Company tender PDF text", text)
            self.assertEqual(parsed.metadata["page_count"], 1)
            self.assertEqual(parsed.metadata["pdf_extraction"], "text")
            self.assertTrue(parsed.units)

    def test_pdf_without_text_reports_clear_error_when_ocr_is_off(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "blank.pdf"
            _make_pdf(path, "")

            with self.assertRaisesRegex(ValueError, "PDF 未提取到文本"):
                parse_file(path, "PDF组", 0, CheckOptions(min_chars=1, pdf_ocr_mode="off"))

    def test_pdf_ocr_reports_progress(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "scan.pdf"
            _make_pdf(path, "")
            messages: list[str] = []

            with patch("checksim.parsers._pdf_text_blocks", return_value=([], {"page_count": 2})), patch(
                "checksim.parsers._paddleocr_pdf_blocks",
                return_value=["第1页: 识别文本"],
            ):
                parsed = parse_file(
                    path,
                    "PDF组",
                    0,
                    CheckOptions(min_chars=1, pdf_ocr_mode="always", pdf_min_text_chars=1),
                    messages.append,
                )

            joined = "\n".join(messages)
            self.assertIn("准备使用 PaddleOCR/PP-OCRv6 识别", joined)
            self.assertIn("共 2 页", joined)
            self.assertIn("OCR 完成", joined)
            self.assertIn("识别出 1 段文本", joined)
            self.assertEqual(parsed.metadata["pdf_extraction"], "paddleocr")

    def test_paddleocr_result_blocks_extracts_rec_texts(self) -> None:
        raw = [{"res": {"rec_texts": ["第一行", "第二行"]}}]
        self.assertEqual(_paddleocr_result_blocks(raw), ["第1页: 第一行", "第1页: 第二行"])

    def test_paddleocr_result_blocks_accepts_legacy_shape(self) -> None:
        raw = [[[[0, 0], [10, 0], [10, 10], [0, 10]], ("Legacy OCR text", 0.98)]]
        self.assertEqual(_paddleocr_result_blocks(raw), ["第1页: Legacy OCR text"])


def _make_docx(path: Path, paragraphs: list[str], image_path: Path | None) -> None:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if image_path is not None:
        doc.add_picture(str(image_path))
    doc.save(path)


def _make_pdf(path: Path, text: str) -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT\n/F1 12 Tf\n72 720 Td\n({escaped}) Tj\nET\n".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"endstream",
    ]
    parts = [b"%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(part) for part in parts))
        parts.append(f"{index} 0 obj\n".encode("ascii") + obj + b"\nendobj\n")
    xref_offset = sum(len(part) for part in parts)
    parts.append(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    parts.append(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        parts.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    parts.append(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(b"".join(parts))


if __name__ == "__main__":
    unittest.main()
