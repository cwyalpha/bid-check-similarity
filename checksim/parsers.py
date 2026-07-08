from __future__ import annotations

import hashlib
import os
import re
import shutil
import sys
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable
from urllib.parse import unquote, urlparse

from docx import Document
from PIL import Image

from .legacy import convert_legacy_to_docx
from .models import CheckOptions, DocumentData, DocumentImage, SUPPORTED_EXTENSIONS, TextUnit, normalize_path
from .text import english_word_count, is_comparable_text, make_ngrams, normalize_for_compare, split_blocks_to_units, visible_char_count


IMAGE_RE = re.compile(r"!\[[^\]]*]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
FENCED_CODE_RE = re.compile(r"```.*?```|~~~.*?~~~", re.DOTALL)
INLINE_CODE_RE = re.compile(r"`[^`]*`")
HTML_TAG_RE = re.compile(r"<[^>]+>")
LINK_RE = re.compile(r"\[([^\]]+)]\(([^)]+)\)")
EMPHASIS_RE = re.compile(r"(\*\*|__)(.*?)\1|(\*|_)([^*_`\n]+)\3|~~(.*?)~~")
ORDERED_LIST_RE = re.compile(r"^\s*\d+[.)、]\s+")
UNORDERED_LIST_RE = re.compile(r"^\s*[-*+]\s+")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


Progress = Callable[[str], None]


def parse_file(
    path: str | Path,
    group_name: str,
    group_index: int,
    options: CheckOptions,
    progress: Progress | None = None,
) -> DocumentData:
    resolved = Path(path).expanduser().resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"文件不存在: {resolved}")
    suffix = resolved.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"暂不支持的文件格式: {resolved.name}")
    if suffix == ".docx":
        return _parse_docx(resolved, group_name, group_index, options)
    if suffix in {".doc", ".wps"}:
        return _parse_legacy_word(resolved, group_name, group_index, options, progress)
    if suffix == ".pdf":
        return _parse_pdf(resolved, group_name, group_index, options, progress)
    if suffix == ".txt":
        return _parse_text(resolved, group_name, group_index, options)
    return _parse_markdown(resolved, group_name, group_index, options)


def _parse_legacy_word(
    path: Path,
    group_name: str,
    group_index: int,
    options: CheckOptions,
    progress: Progress | None,
) -> DocumentData:
    converted_path, temp_dir = convert_legacy_to_docx(
        path,
        log=progress,
        soffice_path=options.soffice_path or None,
        timeout=options.legacy_conversion_timeout,
    )
    try:
        return _parse_docx(converted_path, group_name, group_index, options, source_path=path)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _parse_docx(
    path: Path,
    group_name: str,
    group_index: int,
    options: CheckOptions,
    source_path: Path | None = None,
) -> DocumentData:
    display_path = source_path or path
    document = Document(str(path))
    blocks: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(f"表{table_index} 行{row_index}: " + " | ".join(cells))

    metadata = _docx_metadata(document, display_path)
    if source_path is not None:
        metadata["converted_from"] = source_path.suffix.lower()
        metadata["converted_via"] = "docx"
    images = _docx_images(document, display_path)
    units = _build_units(display_path, group_name, group_index, blocks, options)
    return DocumentData(group_name, group_index, normalize_path(display_path), display_path.name, metadata, blocks, units, images)


def _parse_markdown(path: Path, group_name: str, group_index: int, options: CheckOptions) -> DocumentData:
    raw = _read_text(path)
    images = _markdown_images(raw, path)
    text = _strip_markdown(raw)
    blocks = [line.strip() for line in re.split(r"\n{1,}", text) if line.strip()]
    metadata = _file_metadata(path)
    units = _build_units(path, group_name, group_index, blocks, options)
    return DocumentData(group_name, group_index, normalize_path(path), path.name, metadata, blocks, units, images)


def _parse_text(path: Path, group_name: str, group_index: int, options: CheckOptions) -> DocumentData:
    raw = _read_text(path)
    blocks = [line.strip() for line in re.split(r"\n{1,}", raw) if line.strip()]
    metadata = _file_metadata(path)
    units = _build_units(path, group_name, group_index, blocks, options)
    return DocumentData(group_name, group_index, normalize_path(path), path.name, metadata, blocks, units, [])


def _log(progress: Progress | None, message: str) -> None:
    if progress is not None:
        progress(message)


def _parse_pdf(
    path: Path,
    group_name: str,
    group_index: int,
    options: CheckOptions,
    progress: Progress | None,
) -> DocumentData:
    blocks, metadata = _pdf_text_blocks(path)
    metadata["pdf_extraction"] = "text"
    extracted_chars = visible_char_count("\n".join(blocks))
    if options.pdf_ocr_mode == "off":
        if extracted_chars == 0:
            raise ValueError(f"PDF 未提取到文本: {path.name}。如果这是扫描件，请开启 OCR 或先转换为可复制文本的 PDF。")
        units = _build_units(path, group_name, group_index, blocks, options)
        return DocumentData(group_name, group_index, normalize_path(path), path.name, metadata, blocks, units, [])

    if options.pdf_ocr_mode == "always" or (extracted_chars < options.pdf_min_text_chars and options.pdf_ocr_mode == "auto"):
        page_count = metadata.get("page_count")
        page_hint = f"，共 {page_count} 页" if isinstance(page_count, int) and page_count > 0 else ""
        reason = "强制 OCR" if options.pdf_ocr_mode == "always" else "文本层不足"
        _log(progress, f"  > PDF {reason}，准备使用 PaddleOCR/PP-OCRv6 识别 {path.name}{page_hint}...")
        ocr_started = time.perf_counter()
        blocks = _paddleocr_pdf_blocks(path, options, progress)
        metadata["pdf_extraction"] = "paddleocr"
        extracted_chars = visible_char_count("\n".join(blocks))
        _log(
            progress,
            f"  > OCR 完成: {path.name}，识别出 {len(blocks)} 段文本，用时 {time.perf_counter() - ocr_started:.1f} 秒。",
        )

    if extracted_chars < options.pdf_min_text_chars:
        raise ValueError(
            f"PDF 未提取到足够文本: {path.name}。如果这是扫描件，请安装 PaddleOCR 3.7+ 及其推理引擎后重试，"
            "或在配置中设置 pdf_ocr_mode=\"always\" 强制 OCR。"
        )

    units = _build_units(path, group_name, group_index, blocks, options)
    return DocumentData(group_name, group_index, normalize_path(path), path.name, metadata, blocks, units, [])


def _pdf_text_blocks(path: Path) -> tuple[list[str], dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("读取 PDF 需要安装 pypdf；桌面打包版已内置，源码运行请先安装项目依赖。") from exc

    reader = PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ValueError(f"无法读取加密 PDF: {path.name}") from exc

    blocks: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for line in text.splitlines():
            cleaned = re.sub(r"\s+", " ", line).strip()
            if cleaned:
                blocks.append(f"第{page_index}页: {cleaned}")

    metadata = _file_metadata(path)
    metadata.update(_pdf_metadata(reader, page_count=len(reader.pages)))
    return blocks, metadata


def _pdf_metadata(reader: Any, page_count: int) -> dict[str, Any]:
    raw = reader.metadata or {}

    def value(key: str) -> str:
        item = raw.get(key)
        if item is None:
            return ""
        return str(item)

    return {
        "page_count": page_count,
        "title": value("/Title"),
        "author": value("/Author"),
        "subject": value("/Subject"),
        "creator": value("/Creator"),
        "producer": value("/Producer"),
    }


def _paddleocr_pdf_blocks(path: Path, options: CheckOptions, progress: Progress | None = None) -> list[str]:
    _log(progress, "  > OCR: 正在加载 PaddleOCR 组件...")
    _patch_default_ssl_certs()
    try:
        from paddleocr import PaddleOCR  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "PDF 文本层不足，扫描件 OCR 需要安装 PaddleOCR。请在对应环境执行: "
            "python -m pip install paddleocr，并按 PaddleOCR 文档安装可用推理引擎。"
        ) from exc

    kwargs: dict[str, Any] = {
        "lang": options.pdf_ocr_lang,
        "ocr_version": "PP-OCRv6",
        "use_doc_orientation_classify": False,
        "use_doc_unwarping": False,
        "use_textline_orientation": False,
    }
    if options.pdf_ocr_engine:
        kwargs["engine"] = options.pdf_ocr_engine
    if options.pdf_ocr_det_model:
        kwargs["text_detection_model_name"] = options.pdf_ocr_det_model
    if options.pdf_ocr_rec_model:
        kwargs["text_recognition_model_name"] = options.pdf_ocr_rec_model
    model_root = _ocr_model_root()
    if model_root is not None:
        det_dir = _ocr_model_dir(model_root, options.pdf_ocr_det_model, options.pdf_ocr_engine)
        rec_dir = _ocr_model_dir(model_root, options.pdf_ocr_rec_model, options.pdf_ocr_engine)
        if det_dir is not None:
            kwargs["text_detection_model_dir"] = str(det_dir)
        if rec_dir is not None:
            kwargs["text_recognition_model_dir"] = str(rec_dir)
        if det_dir is not None and rec_dir is not None:
            os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
            _log(progress, "  > OCR: 已找到内置 PP-OCRv6 模型，正在初始化...")
    else:
        _log(progress, "  > OCR: 未找到内置模型目录，将使用 PaddleOCR 默认模型缓存。")
    _log(progress, "  > OCR: 正在初始化模型，首次运行或大文件可能较慢...")
    try:
        try:
            ocr = PaddleOCR(**kwargs)
        except TypeError:
            kwargs.pop("ocr_version", None)
            ocr = PaddleOCR(**kwargs)
    except RuntimeError as exc:
        cause = exc.__cause__
        detail = f"{exc}"
        if cause is not None:
            detail = f"{detail} 原因: {cause}"
        raise RuntimeError(f"PaddleOCR 初始化失败: {detail}") from exc

    _log(progress, f"  > OCR: 模型初始化完成，正在识别 {path.name}...")
    if hasattr(ocr, "predict"):
        raw_results = ocr.predict(str(path))
    else:
        raw_results = ocr.ocr(str(path), cls=False)

    _log(progress, "  > OCR: 识别完成，正在整理文本...")
    blocks = _paddleocr_result_blocks(raw_results)
    if not blocks:
        raise RuntimeError(f"PaddleOCR 未识别到 PDF 文本: {path.name}")
    return blocks


def _patch_default_ssl_certs() -> None:
    try:
        import ssl
    except ImportError:
        return

    try:
        ssl.create_default_context()
        return
    except ssl.SSLError:
        pass

    try:
        import certifi
    except ImportError:
        return

    cafile = certifi.where()

    def load_default_certs(self: ssl.SSLContext, purpose: ssl.Purpose = ssl.Purpose.SERVER_AUTH) -> None:
        self.load_verify_locations(cafile=cafile)

    ssl.SSLContext.load_default_certs = load_default_certs  # type: ignore[method-assign]


def _ocr_model_root() -> Path | None:
    candidates: list[Path] = []
    env_root = os.environ.get("CHECKSIM_OCR_MODEL_DIR")
    if env_root:
        candidates.append(Path(env_root).expanduser())
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", "")
        if meipass:
            candidates.append(Path(meipass) / "ocr_models")
        executable = Path(sys.executable).resolve()
        candidates.append(executable.parent / "ocr_models")
        if len(executable.parents) >= 2:
            candidates.append(executable.parents[1] / "Resources" / "ocr_models")
    candidates.append(Path(__file__).resolve().parents[1] / "packaging" / "ocr_models")

    for candidate in candidates:
        if candidate.is_dir():
            return candidate
    return None


def _ocr_model_dir(root: Path, model_name: str, engine: str) -> Path | None:
    if not model_name:
        return None
    names = [model_name]
    if engine == "onnxruntime" and not model_name.endswith("_onnx"):
        names.insert(0, f"{model_name}_onnx")
    for name in names:
        candidate = root / name
        if (candidate / "inference.onnx").exists() or (candidate / "inference.json").exists():
            return candidate
    return None


def _paddleocr_result_blocks(raw_results: Any) -> list[str]:
    blocks: list[str] = []
    for page_index, result in enumerate(_as_list(raw_results), start=1):
        json_payload = getattr(result, "json", None)
        payload = json_payload() if callable(json_payload) else json_payload if json_payload is not None else result
        texts = _find_ocr_texts(payload)
        for text in texts:
            cleaned = re.sub(r"\s+", " ", str(text)).strip()
            if cleaned:
                blocks.append(f"第{page_index}页: {cleaned}")
    return blocks


def _find_ocr_texts(value: Any) -> list[str]:
    if isinstance(value, dict):
        if "rec_texts" in value and isinstance(value["rec_texts"], (list, tuple)):
            return [str(item) for item in value["rec_texts"] if str(item).strip()]
        texts: list[str] = []
        for item in value.values():
            texts.extend(_find_ocr_texts(item))
        return texts
    if isinstance(value, (list, tuple)):
        if len(value) == 2 and isinstance(value[0], str) and isinstance(value[1], (int, float)):
            return [value[0]] if value[0].strip() else []
        if value and all(isinstance(item, str) for item in value):
            return [item for item in value if item.strip()]
        texts = []
        for item in value:
            texts.extend(_find_ocr_texts(item))
        return texts
    return []


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _build_units(path: Path, group_name: str, group_index: int, blocks: list[str], options: CheckOptions) -> list[TextUnit]:
    units: list[TextUnit] = []
    for index, (location, text) in enumerate(
        split_blocks_to_units(blocks, options.max_unit_chars, options.sentence_delimiters, options.soft_delimiters),
        start=1,
    ):
        normalized = normalize_for_compare(text)
        comparable = is_comparable_text(text, options.min_chars, options.min_words)
        unit_id = _unit_id(path, group_name, location, index, normalized)
        ngrams = make_ngrams(normalized, options.ngram_size) if comparable else tuple()
        units.append(
            TextUnit(
                unit_id=unit_id,
                group_name=group_name,
                group_index=group_index,
                file_path=normalize_path(path),
                file_name=path.name,
                location=location,
                text=text,
                normalized=normalized,
                char_count=visible_char_count(text),
                word_count=english_word_count(text),
                comparable=comparable,
                ngrams=ngrams,
            )
        )
    return units


def _docx_metadata(document: Document, path: Path) -> dict[str, Any]:
    props = document.core_properties
    metadata = _file_metadata(path)
    metadata.update(
        {
            "author": props.author or "",
            "last_modified_by": props.last_modified_by or "",
            "created": _date_value(props.created),
            "modified": _date_value(props.modified),
            "revision": props.revision,
            "title": props.title or "",
            "subject": props.subject or "",
        }
    )
    return metadata


def _file_metadata(path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "path": normalize_path(path),
        "name": path.name,
        "size": stat.st_size,
        "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
    }


def _docx_images(document: Document, path: Path) -> list[DocumentImage]:
    images: list[DocumentImage] = []
    seen: set[str] = set()
    for rel_index, rel in enumerate(document.part.rels.values(), start=1):
        if "image" not in rel.reltype:
            continue
        blob = rel.target_part.blob
        sha = hashlib.sha256(blob).hexdigest()
        if sha in seen:
            continue
        seen.add(sha)
        info = _image_info(blob)
        images.append(
            DocumentImage(
                image_id=f"{path.name}:image{rel_index}",
                source=normalize_path(path),
                sha256=sha,
                ahash=info["ahash"],
                size=len(blob),
                width=info["width"],
                height=info["height"],
            )
        )
    return images


def _markdown_images(raw: str, path: Path) -> list[DocumentImage]:
    images: list[DocumentImage] = []
    for index, match in enumerate(IMAGE_RE.finditer(raw), start=1):
        target = unquote(match.group(1).strip("<>"))
        parsed = urlparse(target)
        if parsed.scheme in {"http", "https", "data"}:
            continue
        image_path = (path.parent / target).resolve()
        if not image_path.exists() or not image_path.is_file():
            continue
        blob = image_path.read_bytes()
        info = _image_info(blob)
        images.append(
            DocumentImage(
                image_id=f"{path.name}:md-image{index}",
                source=normalize_path(image_path),
                sha256=hashlib.sha256(blob).hexdigest(),
                ahash=info["ahash"],
                size=len(blob),
                width=info["width"],
                height=info["height"],
            )
        )
    return images


def _image_info(blob: bytes) -> dict[str, Any]:
    try:
        with Image.open(BytesIO(blob)) as image:
            width, height = image.size
            gray = image.convert("L").resize((8, 8))
            pixels = list(gray.getdata())
            avg = sum(pixels) / len(pixels)
            bits = "".join("1" if pixel >= avg else "0" for pixel in pixels)
            ahash = f"{int(bits, 2):016x}"
            return {"width": width, "height": height, "ahash": ahash}
    except Exception:
        return {"width": None, "height": None, "ahash": None}


def _strip_markdown(raw: str) -> str:
    text = FENCED_CODE_RE.sub("\n", raw)
    text = INLINE_CODE_RE.sub(" ", text)
    text = IMAGE_RE.sub(" ", text)
    text = LINK_RE.sub(r"\1", text)
    text = HTML_TAG_RE.sub(" ", text)
    text = _strip_markdown_emphasis(text)

    lines: list[str] = []
    for line in text.splitlines():
        cleaned = line.strip()
        if not cleaned or TABLE_SEPARATOR_RE.match(cleaned):
            lines.append("")
            continue
        cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", cleaned)
        cleaned = re.sub(r"\s+#{1,6}\s*$", "", cleaned)
        cleaned = re.sub(r"^\s*>\s?", "", cleaned)
        cleaned = ORDERED_LIST_RE.sub("", cleaned)
        cleaned = UNORDERED_LIST_RE.sub("", cleaned)
        cleaned = cleaned.replace("|", " ")
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
        lines.append(cleaned)
    return "\n".join(lines)


def _strip_markdown_emphasis(text: str) -> str:
    def replace(match: re.Match[str]) -> str:
        for group_index in (2, 4, 5):
            value = match.group(group_index)
            if value is not None:
                return value
        return ""

    previous = None
    current = text
    while previous != current:
        previous = current
        current = EMPHASIS_RE.sub(replace, current)
    return current


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _date_value(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat(timespec="seconds")
    return str(value)


def _unit_id(path: Path, group_name: str, location: str, index: int, normalized: str) -> str:
    payload = f"{group_name}|{path}|{location}|{index}|{normalized[:80]}".encode("utf-8", errors="ignore")
    return hashlib.sha1(payload).hexdigest()[:16]
